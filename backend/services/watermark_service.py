"""
Watermark Service for UB360.ai Branding
Adds watermarks to PDF and DOCX files
"""
import os
from pathlib import Path
from typing import Optional
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import settings


class WatermarkService:
    """Add UB360.ai watermarks to documents"""
    
    @staticmethod
    def add_pdf_watermark(input_path: str, output_path: Optional[str] = None) -> str:
        """
        Add UB360.ai watermark to PDF
        
        Args:
            input_path: Path to input PDF
            output_path: Path to output PDF (optional, creates temp if not provided)
        
        Returns:
            Path to watermarked PDF
        """
        if output_path is None:
            output_path = input_path.replace('.pdf', '_watermarked.pdf')
        
        try:
            # Create watermark PDF
            watermark_path = WatermarkService._create_pdf_watermark()
            
            # Read input PDF
            reader = PdfReader(input_path)
            writer = PdfWriter()
            
            # Read watermark
            watermark_reader = PdfReader(watermark_path)
            watermark_page = watermark_reader.pages[0]
            
            # Add watermark to each page
            for page in reader.pages:
                page.merge_page(watermark_page)
                writer.add_page(page)
            
            # Write output
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            # Clean up watermark file
            os.remove(watermark_path)
            
            return output_path
        
        except Exception as e:
            raise Exception(f"Error adding PDF watermark: {str(e)}")
    
    @staticmethod
    def _create_pdf_watermark() -> str:
        """Create watermark PDF overlay"""
        watermark_path = "temp_watermark.pdf"
        c = canvas.Canvas(watermark_path, pagesize=letter)
        width, height = letter
        
        # Set watermark properties
        c.setFont("Helvetica", 10)
        c.setFillColorRGB(0.7, 0.7, 0.7, alpha=0.3)  # Light gray, 30% opacity
        
        # Add diagonal watermark text
        c.saveState()
        c.translate(width/2, height/2)
        c.rotate(45)
        
        # Repeat watermark across page
        for i in range(-2, 3):
            for j in range(-2, 3):
                x = i * 200
                y = j * 150
                c.drawCentredString(x, y, settings.BRAND_WATERMARK)
        
        c.restoreState()
        
        # Add footer
        c.setFillColorRGB(0, 0, 0)
        c.setFont("Helvetica", 8)
        footer_text = f"{settings.BRAND_TAGLINE} | {settings.BRAND_MESSAGE}"
        c.drawCentredString(width/2, 0.5*inch, footer_text)
        
        c.save()
        return watermark_path
    
    @staticmethod
    def add_docx_watermark(input_path: str, output_path: Optional[str] = None) -> str:
        """
        Add UB360.ai watermark to DOCX
        
        Args:
            input_path: Path to input DOCX
            output_path: Path to output DOCX (optional, creates temp if not provided)
        
        Returns:
            Path to watermarked DOCX
        """
        if output_path is None:
            output_path = input_path.replace('.docx', '_watermarked.docx')
        
        try:
            # Open document
            doc = Document(input_path)
            
            # Add header with UB360.ai branding
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = f"{settings.BRAND_NAME} | {settings.BRAND_WATERMARK}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header_para.runs[0]
            header_run.font.size = Pt(9)
            header_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add footer with UB360.ai branding
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = f"{settings.BRAND_TAGLINE} | {settings.BRAND_MESSAGE}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.runs[0]
            footer_run.font.size = Pt(8)
            footer_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Save watermarked document
            doc.save(output_path)
            
            return output_path
        
        except Exception as e:
            raise Exception(f"Error adding DOCX watermark: {str(e)}")
    
    @staticmethod
    def add_watermark(input_path: str, output_path: Optional[str] = None) -> str:
        """
        Add watermark to document (auto-detects type)
        
        Args:
            input_path: Path to input file
            output_path: Path to output file (optional)
        
        Returns:
            Path to watermarked file
        """
        file_ext = Path(input_path).suffix.lower()
        
        if file_ext == '.pdf':
            return WatermarkService.add_pdf_watermark(input_path, output_path)
        elif file_ext == '.docx':
            return WatermarkService.add_docx_watermark(input_path, output_path)
        else:
            # For other file types, just copy the file
            if output_path is None:
                return input_path
            import shutil
            shutil.copy2(input_path, output_path)
            return output_path
    
    @staticmethod
    def format_download_filename(original_filename: str) -> str:
        """
        Format filename with UB360.ai branding
        
        Args:
            original_filename: Original filename
        
        Returns:
            Formatted filename: "{name}..Follow ub360_ai on x.{ext}"
        """
        # Split name and extension
        path = Path(original_filename)
        name = path.stem
        ext = path.suffix
        
        # Add branding suffix
        branded_name = f"{name}{settings.BRAND_FILENAME_SUFFIX}{ext}"
        
        return branded_name
