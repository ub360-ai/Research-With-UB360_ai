"""
DOCX document handler
Extracts text and metadata from Word documents
"""
from docx import Document
from typing import Dict, Any, List
from datetime import datetime


class DOCXHandler:
    """Handle DOCX document processing"""
    
    @staticmethod
    def extract_text_and_metadata(file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX file
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Dictionary with text and metadata
        """
        result = {
            "text": "",
            "paragraphs": [],
            "metadata": {
                "total_paragraphs": 0,
                "author": None,
                "title": None,
                "subject": None,
                "keywords": None,
                "created": None,
                "modified": None,
                "last_modified_by": None
            }
        }
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs_text.append(para.text)
                    result["paragraphs"].append({
                        "text": para.text,
                        "style": para.style.name if para.style else None
                    })
            
            result["text"] = "\n\n".join(paragraphs_text)
            result["metadata"]["total_paragraphs"] = len(paragraphs_text)
            
            # Extract metadata from core properties
            core_props = doc.core_properties
            
            result["metadata"]["author"] = core_props.author
            result["metadata"]["title"] = core_props.title
            result["metadata"]["subject"] = core_props.subject
            result["metadata"]["keywords"] = core_props.keywords
            result["metadata"]["last_modified_by"] = core_props.last_modified_by
            
            # Handle dates
            if core_props.created:
                result["metadata"]["created"] = core_props.created.isoformat()
            if core_props.modified:
                result["metadata"]["modified"] = core_props.modified.isoformat()
            
            # Filter out None values from metadata (ChromaDB doesn't accept None)
            result["metadata"] = {
                k: v for k, v in result["metadata"].items() 
                if v is not None
            }
            
            return result
        
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_simple(file_path: str) -> str:
        """
        Simple text extraction
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Extracted text
        """
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    @staticmethod
    def extract_tables(file_path: str) -> List[List[List[str]]]:
        """
        Extract tables from DOCX
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            List of tables (each table is a list of rows)
        """
        try:
            doc = Document(file_path)
            tables_data = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            return tables_data
        except Exception as e:
            raise Exception(f"Error extracting tables: {str(e)}")
