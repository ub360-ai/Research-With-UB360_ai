"""
DOCX report generator using python-docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import List, Dict, Any
import io


class DOCXGenerator:
    """Generate DOCX reports for research findings"""
    
    @staticmethod
    def generate_research_report(
        title: str,
        queries: List[Dict[str, Any]],
        output_path: str = None
    ) -> bytes:
        """
        Generate a research report DOCX
        
        Args:
            title: Report title
            queries: List of query results with answers and citations
            output_path: Optional file path to save DOCX
        
        Returns:
            DOCX content as bytes
        """
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = "Research With UB360.ai"
        doc.core_properties.created = datetime.now()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Add queries and answers
        for idx, query_data in enumerate(queries, 1):
            # Question
            question_heading = doc.add_heading(f"Question {idx}", level=2)
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(query_data.get('question', 'N/A'))
            question_run.bold = True
            
            # Answer
            doc.add_heading("Answer:", level=3)
            answer_para = doc.add_paragraph(query_data.get('answer', 'No answer provided'))
            answer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Citations
            citations = query_data.get('citations', [])
            if citations:
                doc.add_heading("Sources:", level=3)
                for cite_idx, citation in enumerate(citations[:5], 1):
                    doc_name = citation.get('document_name', 'Unknown')
                    page_num = citation.get('page_number')
                    score = citation.get('relevance_score', 0)
                    
                    cite_text = f"{cite_idx}. {doc_name}"
                    if page_num:
                        cite_text += f" (Page {page_num})"
                    cite_text += f" - Relevance: {score:.2%}"
                    
                    cite_para = doc.add_paragraph(cite_text, style='List Number')
                    cite_para.paragraph_format.left_indent = Inches(0.5)
            
            # Add spacing between questions
            doc.add_paragraph()
        
        # Save or return bytes
        if output_path:
            doc.save(output_path)
            return None
        else:
            # Save to BytesIO
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            return file_stream.getvalue()
    
    @staticmethod
    def generate_bibliography(
        documents: List[Dict[str, Any]],
        style: str = "APA",
        output_path: str = None
    ) -> bytes:
        """
        Generate a bibliography DOCX
        
        Args:
            documents: List of document metadata
            style: Citation style (APA, MLA, Chicago)
            output_path: Optional file path to save DOCX
        
        Returns:
            DOCX content as bytes
        """
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = f"Bibliography ({style} Style)"
        doc.core_properties.author = "Research With UB360.ai"
        
        # Add title
        title_para = doc.add_heading(f"Bibliography ({style} Style)", 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Sort documents alphabetically
        sorted_docs = sorted(documents, key=lambda x: x.get('filename', ''))
        
        # Add bibliography entries
        for doc_data in sorted_docs:
            citation_text = DOCXGenerator._format_citation(doc_data, style)
            para = doc.add_paragraph(citation_text)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)
        
        # Save or return bytes
        if output_path:
            doc.save(output_path)
            return None
        else:
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            return file_stream.getvalue()
    
    @staticmethod
    def generate_summary(
        title: str,
        summary_text: str,
        sources: List[Dict[str, Any]],
        output_path: str = None
    ) -> bytes:
        """
        Generate a research summary DOCX
        
        Args:
            title: Summary title
            summary_text: Main summary content
            sources: List of source documents
            output_path: Optional file path to save DOCX
        
        Returns:
            DOCX content as bytes
        """
        # Create document
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.add_run(datetime.now().strftime('%B %d, %Y'))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Add summary
        doc.add_heading("Summary", level=1)
        summary_para = doc.add_paragraph(summary_text)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        
        # Add sources
        if sources:
            doc.add_heading("Sources Consulted", level=1)
            for idx, source in enumerate(sources, 1):
                doc_name = source.get('document_name', 'Unknown')
                para = doc.add_paragraph(f"{idx}. {doc_name}", style='List Number')
        
        # Save or return bytes
        if output_path:
            doc.save(output_path)
            return None
        else:
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            return file_stream.getvalue()
    
    @staticmethod
    def _format_citation(doc: Dict[str, Any], style: str) -> str:
        """Format a single citation"""
        metadata = doc.get('metadata', {})
        filename = doc.get('filename', 'Unknown')
        doc_type = doc.get('document_type', 'unknown')
        
        author = metadata.get('author', 'Unknown Author')
        title = metadata.get('title', filename)
        year = 'n.d.'
        
        # Extract year
        if 'creation_date' in metadata:
            try:
                year = metadata['creation_date'][:4]
            except:
                pass
        elif 'created' in metadata:
            try:
                year = metadata['created'][:4]
            except:
                pass
        
        # Format based on style
        if style.upper() == "APA":
            citation = f"{author}. ({year}). {title}."
            if doc_type == "url":
                url = metadata.get('url', '')
                citation += f" Retrieved from {url}"
        
        elif style.upper() == "MLA":
            citation = f"{author}. \"{title}.\" {year}."
        
        elif style.upper() == "CHICAGO":
            citation = f"{author}. {title}. {year}."
        
        else:
            citation = f"{author}. {title}. {year}."
        
        return citation
