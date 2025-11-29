"""
Chat History Exporter with UB360.ai Branding and Professional Formatting
Exports conversation history in PDF, DOCX, or JSON format with proper spacing and structure
"""
import json
import re
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any
from io import BytesIO

# PDF dependencies
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Preformatted, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# DOCX dependencies
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Markdown processing
import markdown2
from bs4 import BeautifulSoup

from config import settings


class ChatExporter:
    """Export chat conversations with UB360.ai branding and professional formatting"""
    
    @staticmethod
    def export_conversation(
        conversation_title: str,
        messages: List[Dict[str, Any]],
        format: str = 'pdf'
    ) -> bytes:
        """
        Export a single conversation
        
        Args:
            conversation_title: Title of the conversation
            messages: List of message dicts with 'role' and 'content'
            format: Export format ('pdf', 'docx', 'json')
        
        Returns:
            Exported file as bytes
        """
        if format == 'pdf':
            return ChatExporter._export_pdf(conversation_title, messages)
        elif format == 'docx':
            return ChatExporter._export_docx(conversation_title, messages)
        elif format == 'json':
            return ChatExporter._export_json(conversation_title, messages)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    @staticmethod
    def _escape_for_reportlab(text: str) -> str:
        """
        Escape special characters for ReportLab XML
        
        Args:
            text: Text to escape
        
        Returns:
            Escaped text safe for ReportLab
        """
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        return text
    
    @staticmethod
    def _split_code_block(code: str, max_lines: int = 50) -> List[str]:
        """
        Split large code blocks into manageable chunks
        
        Args:
            code: Code text to split
            max_lines: Maximum lines per chunk
        
        Returns:
            List of code chunks
        """
        lines = code.split('\n')
        chunks = []
        
        for i in range(0, len(lines), max_lines):
            chunk = '\n'.join(lines[i:i + max_lines])
            chunks.append(chunk)
        
        return chunks
    
    @staticmethod
    def _wrap_long_lines(code: str, max_length: int = 80) -> str:
        """
        Wrap very long lines in code to prevent overflow
        
        Args:
            code: Code text
            max_length: Maximum line length
        
        Returns:
            Code with wrapped lines
        """
        lines = code.split('\n')
        wrapped_lines = []
        
        for line in lines:
            if len(line) <= max_length:
                wrapped_lines.append(line)
            else:
                # Wrap at max_length, preserving indentation
                indent = len(line) - len(line.lstrip())
                indent_str = ' ' * indent
                
                remaining = line
                while len(remaining) > max_length:
                    wrapped_lines.append(remaining[:max_length])
                    remaining = indent_str + '  ' + remaining[max_length:]  # Add extra indent for continuation
                
                if remaining.strip():
                    wrapped_lines.append(remaining)
        
        return '\n'.join(wrapped_lines)
    
    @staticmethod
    def _split_into_paragraphs(text: str) -> List[str]:
        """Split text into paragraphs for better readability"""
        # Split on double newlines or single newlines followed by bullet points
        paragraphs = re.split(r'\n\n+|\n(?=[•\-\*\d+\.])', text)
        return [p.strip() for p in paragraphs if p.strip()]
    
    @staticmethod
    def _markdown_to_reportlab(markdown_text: str) -> str:
        """
        Convert markdown to ReportLab XML tags with proper formatting
        
        Args:
            markdown_text: Markdown formatted text
        
        Returns:
            Text with ReportLab XML tags
        """
        # Convert markdown to HTML
        html = markdown2.markdown(
            markdown_text,
            extras=['fenced-code-blocks', 'tables', 'break-on-newline', 'code-friendly']
        )
        
        # Parse HTML
        soup = BeautifulSoup(html, 'html.parser')
        
        # Convert HTML to ReportLab XML
        text = markdown_text
        
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'__(.+?)__', r'<b>\1</b>', text)
        
        # Italic
        text = re.sub(r'\*([^*]+?)\*', r'<i>\1</i>', text)
        text = re.sub(r'_([^_]+?)_', r'<i>\1</i>', text)
        
        # Inline code
        text = re.sub(r'`([^`]+?)`', r"<font name='Courier' color='#d63031' size='10'>\1</font>", text)
        
        # Convert newlines to line breaks
        text = text.replace('\n', '<br/>')
        
        return text
    
    @staticmethod
    def _export_pdf(title: str, messages: List[Dict]) -> bytes:
        """Export conversation as PDF with professional spacing and formatting"""
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=1*inch,
                bottomMargin=1*inch
            )
            
            # Build content
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles with better spacing
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#10A37F'),
                spaceAfter=40,
                spaceBefore=10,
                alignment=TA_CENTER,
                leading=28
            )
            
            brand_style = ParagraphStyle(
                'Brand',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.grey,
                alignment=TA_CENTER,
                spaceAfter=30,
                leading=14
            )
            
            user_label_style = ParagraphStyle(
                'UserLabel',
                parent=styles['Normal'],
                fontSize=12,
                fontName='Helvetica-Bold',
                textColor=colors.HexColor('#2C3E50'),
                spaceAfter=8,
                spaceBefore=20,
                alignment=TA_LEFT
            )
            
            user_style = ParagraphStyle(
                'UserMessage',
                parent=styles['Normal'],
                fontSize=11,
                alignment=TA_LEFT,
                textColor=colors.HexColor('#2C3E50'),
                spaceAfter=20,
                leftIndent=20,
                rightIndent=20,
                leading=16,
                backColor=colors.HexColor('#F8F9FA')
            )
            
            ai_label_style = ParagraphStyle(
                'AILabel',
                parent=styles['Normal'],
                fontSize=12,
                fontName='Helvetica-Bold',
                textColor=colors.HexColor('#10A37F'),
                spaceAfter=8,
                spaceBefore=20,
                alignment=TA_LEFT
            )
            
            ai_style = ParagraphStyle(
                'AIMessage',
                parent=styles['Normal'],
                fontSize=11,
                alignment=TA_LEFT,
                textColor=colors.HexColor('#000000'),
                spaceAfter=20,
                leftIndent=20,
                rightIndent=20,
                leading=16
            )
            
            code_style = ParagraphStyle(
                'CodeBlock',
                parent=styles['Code'],
                fontSize=9,
                fontName='Courier',
                textColor=colors.HexColor('#2C3E50'),
                backColor=colors.HexColor('#F5F5F5'),
                leftIndent=30,
                rightIndent=30,
                spaceAfter=15,
                spaceBefore=10,
                leading=12,
                borderPadding=10
            )
            
            # Header with branding
            story.append(Paragraph(f"<b>{title}</b>", title_style))
            story.append(Paragraph(
                f"<i>{settings.BRAND_TAGLINE}</i>",
                brand_style
            ))
            story.append(Paragraph(
                f"<i>{settings.BRAND_MESSAGE}</i>",
                brand_style
            ))
            story.append(Spacer(1, 0.4*inch))
            
            # Messages with better spacing
            for i, msg in enumerate(messages):
                role = msg.get('role', msg.get('type', 'user'))
                content = msg.get('content', '')
                
                # Check for code blocks
                if '```' in content:
                    # Handle code blocks separately
                    parts = content.split('```')
                    
                    # Add label
                    if role in ['user', 'human']:
                        story.append(Paragraph("<b>You:</b>", user_label_style))
                    else:
                        story.append(Paragraph("<b>Professor UB360:</b>", ai_label_style))
                    
                    for idx, part in enumerate(parts):
                        if idx % 2 == 0 and part.strip():  # Regular text
                            # Split into paragraphs
                            paragraphs = ChatExporter._split_into_paragraphs(part)
                            for para in paragraphs:
                                formatted_text = ChatExporter._markdown_to_reportlab(para)
                                style = user_style if role in ['user', 'human'] else ai_style
                                story.append(Paragraph(formatted_text, style))
                                story.append(Spacer(1, 0.1*inch))
                        elif idx % 2 == 1:  # Code block
                            # Remove language identifier if present
                            code_text = re.sub(r'^[a-z]+\n', '', part, flags=re.IGNORECASE)
                            code_text = code_text.strip()
                            
                            if not code_text:
                                continue
                            
                            # Escape special characters for ReportLab
                            code_text = ChatExporter._escape_for_reportlab(code_text)
                            
                            # Check if code is too large or has very long lines
                            lines = code_text.split('\n')
                            is_large = len(lines) > 50 or any(len(line) > 100 for line in lines)
                            
                            if is_large:
                                # Split into manageable chunks
                                chunks = ChatExporter._split_code_block(code_text, max_lines=50)
                                
                                for chunk_idx, chunk in enumerate(chunks):
                                    try:
                                        # Wrap long lines
                                        chunk = ChatExporter._wrap_long_lines(chunk, max_length=80)
                                        story.append(Preformatted(chunk, code_style))
                                        
                                        # Add continuation indicator if not last chunk
                                        if chunk_idx < len(chunks) - 1:
                                            cont_style = user_style if role in ['user', 'human'] else ai_style
                                            story.append(Paragraph("<i>...continued...</i>", cont_style))
                                            story.append(Spacer(1, 0.05*inch))
                                    except Exception as e:
                                        print(f"Error adding code chunk {chunk_idx}: {e}")
                                        # Fallback: add as regular paragraph
                                        fallback_style = user_style if role in ['user', 'human'] else ai_style
                                        story.append(Paragraph(f"<font name='Courier' size='9'>{chunk[:500]}...</font>", fallback_style))
                            else:
                                # Small code block - add directly
                                try:
                                    story.append(Preformatted(code_text, code_style))
                                except Exception as e:
                                    print(f"Error adding code block: {e}")
                                    # Fallback: add as regular paragraph
                                    fallback_style = user_style if role in ['user', 'human'] else ai_style
                                    story.append(Paragraph(f"<font name='Courier' size='9'>{code_text[:500]}...</font>", fallback_style))
                            
                            story.append(Spacer(1, 0.1*inch))
                else:
                    # No code blocks - process as regular text with paragraphs
                    if role in ['user', 'human']:
                        story.append(Paragraph("<b>You:</b>", user_label_style))
                    else:
                        story.append(Paragraph("<b>Professor UB360:</b>", ai_label_style))
                    
                    # Split into paragraphs for better readability
                    paragraphs = ChatExporter._split_into_paragraphs(content)
                    
                    for para in paragraphs:
                        formatted_text = ChatExporter._markdown_to_reportlab(para)
                        style = user_style if role in ['user', 'human'] else ai_style
                        story.append(Paragraph(formatted_text, style))
                        story.append(Spacer(1, 0.1*inch))
                
                # Add extra space between messages
                story.append(Spacer(1, 0.2*inch))
            
            # Footer watermark
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph(
                f"<i>{settings.BRAND_WATERMARK}</i>",
                brand_style
            ))
            
            # Build PDF
            doc.build(story, onFirstPage=ChatExporter._add_pdf_header_footer,
                      onLaterPages=ChatExporter._add_pdf_header_footer)
            
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            print(f"Error exporting PDF: {e}")
            import traceback
            traceback.print_exc()
            
            # Return minimal PDF with error message
            return ChatExporter._create_error_pdf(title, str(e))
    
    @staticmethod
    def _create_error_pdf(title: str, error: str) -> bytes:
        """Create a simple PDF with error message when export fails"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        
        story = []
        styles = getSampleStyleSheet()
        
        story.append(Paragraph(f"<b>Export Error</b>", styles['Title']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f"Failed to export: {title}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"<b>Error:</b> {error}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Please try exporting a smaller conversation or contact support.", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(f"<i>{settings.BRAND_WATERMARK}</i>", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def _add_pdf_header_footer(canvas, doc):
        """Add header and footer to PDF pages"""
        canvas.saveState()
        width, height = letter
        
        # Header
        canvas.setFont('Helvetica-Bold', 10)
        canvas.setFillColorRGB(0.06, 0.64, 0.50)  # #10A37F
        canvas.drawCentredString(width/2, height - 0.5*inch, "Research with UB360.ai")
        
        # Footer
        canvas.setFont('Helvetica', 8)
        canvas.setFillGray(0.5)
        canvas.drawCentredString(
            width/2,
            0.5*inch,
            f"{settings.BRAND_WATERMARK} | Page {doc.page}"
        )
        
        canvas.restoreState()
    
    @staticmethod
    def _add_formatted_text_to_docx(paragraph, text: str, is_code: bool = False):
        """
        Add formatted text to DOCX paragraph with proper markdown rendering
        
        Args:
            paragraph: DOCX paragraph object
            text: Text to add (may contain markdown)
            is_code: Whether this is a code block
        """
        if is_code:
            # Code block - add as monospace
            run = paragraph.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(44, 62, 80)
            return
        
        # Parse markdown inline formatting
        # Split by markdown patterns while preserving them
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        
        for part in parts:
            if not part:
                continue
                
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
                run.font.size = Pt(11)
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True
                run.font.size = Pt(11)
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(214, 48, 49)
            else:
                # Regular text
                run = paragraph.add_run(part)
                run.font.size = Pt(11)
    
    @staticmethod
    def _export_docx(title: str, messages: List[Dict]) -> bytes:
        """Export conversation as DOCX with professional spacing and formatting"""
        doc = Document()
        
        # Add header
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = f"{settings.BRAND_NAME} | {settings.BRAND_WATERMARK}"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.runs[0]
        header_run.font.size = Pt(9)
        header_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = f"{settings.BRAND_TAGLINE} | {settings.BRAND_MESSAGE}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(16, 163, 127)  # #10A37F
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(20)
        
        # Branding
        brand_para = doc.add_paragraph()
        brand_run = brand_para.add_run(f"{settings.BRAND_TAGLINE}\n{settings.BRAND_MESSAGE}")
        brand_run.font.size = Pt(10)
        brand_run.font.italic = True
        brand_run.font.color.rgb = RGBColor(128, 128, 128)
        brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        brand_para.space_after = Pt(30)
        
        # Messages with better spacing
        for msg in messages:
            role = msg.get('role', msg.get('type', 'user'))
            content = msg.get('content', '')
            
            # Add label with spacing
            label_para = doc.add_paragraph()
            label_para.space_before = Pt(20)
            label_para.space_after = Pt(8)
            
            if role in ['user', 'human']:
                label_run = label_para.add_run("You:")
                label_run.font.bold = True
                label_run.font.size = Pt(12)
                label_run.font.color.rgb = RGBColor(44, 62, 80)
            else:
                label_run = label_para.add_run("Professor UB360:")
                label_run.font.bold = True
                label_run.font.size = Pt(12)
                label_run.font.color.rgb = RGBColor(16, 163, 127)
            
            # Handle code blocks
            if '```' in content:
                parts = content.split('```')
                for idx, part in enumerate(parts):
                    if idx % 2 == 0 and part.strip():  # Regular text
                        paragraphs = ChatExporter._split_into_paragraphs(part)
                        for para_text in paragraphs:
                            msg_para = doc.add_paragraph()
                            msg_para.paragraph_format.left_indent = Inches(0.3)
                            msg_para.paragraph_format.right_indent = Inches(0.3)
                            msg_para.space_after = Pt(10)
                            ChatExporter._add_formatted_text_to_docx(msg_para, para_text)
                    elif idx % 2 == 1:  # Code block
                        # Remove language identifier
                        code_text = re.sub(r'^[a-z]+\n', '', part, flags=re.IGNORECASE)
                        code_text = code_text.strip()
                        
                        if not code_text:
                            continue
                        
                        # Check if code is very large
                        lines = code_text.split('\n')
                        
                        if len(lines) > 100:
                            # Split into chunks of 100 lines
                            chunks = ChatExporter._split_code_block(code_text, max_lines=100)
                            
                            for chunk_idx, chunk in enumerate(chunks):
                                code_para = doc.add_paragraph()
                                code_para.paragraph_format.left_indent = Inches(0.5)
                                code_para.paragraph_format.right_indent = Inches(0.5)
                                code_para.space_after = Pt(5)
                                code_para.space_before = Pt(5)
                                
                                ChatExporter._add_formatted_text_to_docx(code_para, chunk, is_code=True)
                                
                                # Add continuation indicator
                                if chunk_idx < len(chunks) - 1:
                                    cont_para = doc.add_paragraph()
                                    cont_run = cont_para.add_run("...continued...")
                                    cont_run.font.italic = True
                                    cont_run.font.size = Pt(9)
                                    cont_run.font.color.rgb = RGBColor(128, 128, 128)
                                    cont_para.space_after = Pt(5)
                        else:
                            # Normal size code block
                            code_para = doc.add_paragraph()
                            code_para.paragraph_format.left_indent = Inches(0.5)
                            code_para.paragraph_format.right_indent = Inches(0.5)
                            code_para.space_after = Pt(15)
                            code_para.space_before = Pt(10)
                            
                            ChatExporter._add_formatted_text_to_docx(code_para, code_text, is_code=True)
            else:
                # Split into paragraphs for better readability
                paragraphs = ChatExporter._split_into_paragraphs(content)
                
                for para_text in paragraphs:
                    msg_para = doc.add_paragraph()
                    msg_para.paragraph_format.left_indent = Inches(0.3)
                    msg_para.paragraph_format.right_indent = Inches(0.3)
                    msg_para.space_after = Pt(10)
                    msg_para.paragraph_format.line_spacing = 1.3
                    
                    ChatExporter._add_formatted_text_to_docx(msg_para, para_text)
            
            # Extra space between messages
            spacer = doc.add_paragraph()
            spacer.space_after = Pt(10)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def _export_json(title: str, messages: List[Dict]) -> bytes:
        """Export conversation as JSON"""
        export_data = {
            "title": title,
            "exported_at": datetime.now().isoformat(),
            "exported_by": settings.BRAND_NAME,
            "follow_us": settings.BRAND_HANDLE,
            "platform": settings.BRAND_PLATFORM,
            "messages": messages,
            "metadata": {
                "total_messages": len(messages),
                "branding": settings.BRAND_TAGLINE
            }
        }
        
        json_str = json.dumps(export_data, indent=2, ensure_ascii=False)
        return json_str.encode('utf-8')
    
    @staticmethod
    def format_export_filename(title: str, format: str) -> str:
        """
        Format export filename with UB360.ai branding
        
        Args:
            title: Conversation title
            format: File format (pdf, docx, json)
        
        Returns:
            Formatted filename
        """
        # Clean title
        clean_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_'))
        clean_title = clean_title.strip()[:50]  # Limit length
        
        # Add branding suffix
        branded_name = f"{clean_title}{settings.BRAND_FILENAME_SUFFIX}.{format}"
        
        return branded_name
